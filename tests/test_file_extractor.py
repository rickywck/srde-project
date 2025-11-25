"""
Tests for file_extractor module
Tests text extraction from various file formats
"""

import pytest
from io import BytesIO
from pathlib import Path
import tempfile
import os

from tools.file_extractor import FileExtractor, extract_text_from_file


class TestFileExtractor:
    """Test suite for FileExtractor class"""
    
    def test_extract_text_from_txt_bytes(self):
        """Test extraction from plain text file (bytes)"""
        content = b"Hello, World!\nThis is a test."
        result = FileExtractor.extract_text(content, "test.txt")
        assert result == "Hello, World!\nThis is a test."
    
    def test_extract_text_from_txt_bytesio(self):
        """Test extraction from plain text file (BytesIO)"""
        content = BytesIO(b"Hello, World!\nThis is a test.")
        result = FileExtractor.extract_text(content, "test.txt")
        assert result == "Hello, World!\nThis is a test."
    
    def test_extract_text_from_md(self):
        """Test extraction from markdown file"""
        content = b"# Header\n\nThis is **markdown** content."
        result = FileExtractor.extract_text(content, "test.md")
        assert "# Header" in result
        assert "markdown" in result
    
    def test_extract_text_from_docx(self):
        """Test extraction from DOCX file"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        
        # Create a simple DOCX file in memory
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        
        # Save to BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        result = FileExtractor.extract_text(docx_bytes, "test.docx")
        assert "First paragraph" in result
        assert "Second paragraph" in result
    
    def test_extract_text_from_pdf(self):
        """Test extraction from PDF file"""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            pytest.skip("PyPDF2 not installed")
        
        # Create a simple PDF for testing
        # Note: This test requires a sample PDF file
        # For a real test, you'd need to create or include a sample PDF
        pytest.skip("PDF creation requires additional libraries (reportlab). Manual test recommended.")
    
    def test_extract_text_from_file_path(self):
        """Test extraction from file path"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Test content from file")
            temp_path = f.name
        
        try:
            result = FileExtractor.extract_text(temp_path, os.path.basename(temp_path))
            assert "Test content from file" in result
        finally:
            os.unlink(temp_path)
    
    def test_extract_text_from_path_object(self):
        """Test extraction from Path object"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("Test content from Path")
            temp_path = Path(f.name)
        
        try:
            result = FileExtractor.extract_text(temp_path, temp_path.name)
            assert "Test content from Path" in result
        finally:
            temp_path.unlink()
    
    def test_unsupported_format(self):
        """Test handling of unsupported file format"""
        content = b"\x00\x01\x02\x03"  # Binary content
        with pytest.raises(ValueError, match="Unsupported file format"):
            FileExtractor.extract_text(content, "test.bin")
    
    def test_is_supported(self):
        """Test is_supported method"""
        assert FileExtractor.is_supported("test.txt")
        assert FileExtractor.is_supported("test.md")
        assert FileExtractor.is_supported("test.docx")
        assert FileExtractor.is_supported("test.pdf")
        assert FileExtractor.is_supported("TEST.TXT")  # Case insensitive
        assert not FileExtractor.is_supported("test.bin")
        assert not FileExtractor.is_supported("test.jpg")
    
    def test_get_supported_extensions(self):
        """Test get_supported_extensions method"""
        extensions = FileExtractor.get_supported_extensions()
        assert '.txt' in extensions
        assert '.md' in extensions
        assert '.docx' in extensions
        assert '.pdf' in extensions
        assert len(extensions) == 4
    
    def test_convenience_function(self):
        """Test convenience function extract_text_from_file"""
        content = b"Testing convenience function"
        result = extract_text_from_file(content, "test.txt")
        assert result == "Testing convenience function"
    
    def test_utf8_encoding(self):
        """Test UTF-8 encoding handling"""
        content = "Hello 世界 🌍".encode('utf-8')
        result = FileExtractor.extract_text(content, "test.txt")
        assert "Hello" in result
        assert "世界" in result
        assert "🌍" in result
    
    def test_empty_file(self):
        """Test extraction from empty file"""
        content = b""
        result = FileExtractor.extract_text(content, "empty.txt")
        assert result == ""
    
    def test_docx_with_empty_paragraphs(self):
        """Test DOCX extraction with empty paragraphs"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        
        doc = Document()
        doc.add_paragraph("First")
        doc.add_paragraph("")  # Empty paragraph
        doc.add_paragraph("Second")
        
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        result = FileExtractor.extract_text(docx_bytes, "test.docx")
        assert "First" in result
        assert "Second" in result


if __name__ == "__main__":
    # Run tests
    pytest.main([__file__, "-v"])
