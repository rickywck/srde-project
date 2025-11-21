import os
from pinecone import Pinecone
from docx import Document
import PyPDF2  # Library for extracting text from PDFs
# Main Chunking Functions
from chunking_evaluation.chunking import (
    RecursiveTokenChunker
) 

# Prompt user for input parameters
file_path = input("Enter the path to the file (.docx or .pdf): ").strip()
namespace = input("Enter the namespace for the Pinecone index (e.g., gentingbrd-namespace): ").strip()
chunk_size = input("Enter the chunk size (default is 500): ").strip()
overlap_size = input("Enter the overlap size (default is 100): ").strip()

# Set default chunk size and overlap size if not provided
chunk_size = int(chunk_size) if chunk_size.isdigit() else 500
overlap_size = int(overlap_size) if overlap_size.isdigit() else 100

# Initialize Pinecone
api_key = os.getenv("PINECONE_API_KEY")
if not api_key:
    raise ValueError("PINECONE_API_KEY environment variable is not set")

pc = Pinecone(api_key=api_key)

# Target the index where you'll store the vector embeddings
index = pc.Index("brd-upload")

# Function to split text into chunks with overlap
def split_text_into_chunks_with_overlap(text, chunk_size=500, overlap_size=100):
    words = text.split()
    start = 0
    while start < len(words):
        end = start + chunk_size
        yield " ".join(words[start:end])
        start += chunk_size - overlap_size  # Move the start forward with overlap

# helper function to split text into chunks with recursive method
recursive_character_chunker = RecursiveTokenChunker(
    chunk_size=chunk_size*2,  # Character Length
    chunk_overlap=overlap_size,  # Overlap
    length_function=len,  # Character length with len()
    separators=["\n\n", "\n", ".", "?", "!", " ", ""] # According to Research
)

# Function to extract text from a .docx file
def extract_text_from_docx(file_path):
    document = Document(file_path)
    paragraphs = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():  # Skip empty paragraphs
            paragraphs.append(paragraph.text.strip())
    table_texts = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Skip empty cells
                    table_texts.append(cell.text.strip())
    return "\n".join(paragraphs + table_texts)

# Function to extract text from a .pdf file
def extract_text_from_pdf(file_path):
    pdf_text = []
    with open(file_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():  # Skip empty pages
                pdf_text.append(text.strip())
    return "\n".join(pdf_text)

# Determine the file type and extract text
if not os.path.exists(file_path):
    raise FileNotFoundError(f"The file '{file_path}' does not exist.")

file_extension = os.path.splitext(file_path)[1].lower()
if file_extension == ".docx":
    all_text = extract_text_from_docx(file_path)
elif file_extension == ".pdf":
    all_text = extract_text_from_pdf(file_path)
else:
    raise ValueError("Unsupported file type. Please provide a .docx or .pdf file.")

# Debug: Print the size and content of the extracted text
print("Size of extracted text:", len(all_text))
print("Extracted text:", all_text)


# Split the text into chunks with overlap
#chunks = list(split_text_into_chunks_with_overlap(all_text, chunk_size, overlap_size))

chunks = recursive_character_chunker.split_text(all_text)

# Transform chunks into an array of dictionaries
records = [{"id": f"vec{i+1}", "text": chunk} for i, chunk in enumerate(chunks)]

# Print the size of the records
print(f"Number of records: {len(chunks)}")

# Insert the records into the index
index.upsert_records(namespace, records)
